import json, os
from uuid import uuid4
import docx
import datetime


template = {
    "id": None,
    "last_names": [],
    "first_names": [],
    "date_of_birth": None,
    "place_of_birth": None,
    "date_of_death": None,
    "place_of_death": None,
    "direct_relatives": {"ascendants": [], "descendants": [], "partners": []},
    "life_events": [
        {
            "start": str(),
            "finish": None,
            "description": "Lifetime",
            "notes": None,
            "sources": [],
        }
    ],
}


class person:
    """A person record"""

    def __init__(self):
        self.data = template
        self.filename = str()

    def create(self, **kwargs):
        self.data["id"] = str(uuid4().hex)[:6]
        self.data["last_names"] = kwargs.get("ln", [])
        self.data["first_names"] = kwargs.get("fn", [])
        self.data["date_of_birth"] = kwargs.get("dob", None)
        self.data["place_of_birth"] = kwargs.get("pob", None)
        self.filename = f"{self.data['id']}.json"

        if self.data["date_of_birth"]:
            self.data["life_events"][0]["start"] = self.data["date_of_birth"]

    def load(self, id):
        self.filename = f"{id}.json"
        assert os.access(
            self.filename, os.F_OK
        ), f"Filename {self.filename} does not exists"
        self.data = json.load(open(self.filename, "r", encoding="utf-8"))

    def save(self):
        with open(self.filename, "w", encoding="utf-8") as file:
            json.dump(self.data, file, indent=4)
            print(f"{self.filename} saved.")

    def print(self):
        print(json.dumps(self.data, indent=4))

    def add_event(self, **kwargs):
        self.data["life_events"].append(
            {
                "start": kwargs.get("start", str()),
                "finish": kwargs.get("finish", None),
                "description": kwargs.get("description", None),
                "notes": kwargs.get("notes", None),
                "sources": kwargs.get("sources", []),
            }
        )

    def sort_life_events(self):

        self.data["life_events"].sort(key=lambda x: dict.get(x, "start"))

    def print_unsourced_events(self):
        for event in self.data["life_events"]:
            if len(event["sources"]) == 0:
                print(f"{event['start']} {event['description']}")


def link_child(**kwargs):
    """Add child as descendants and parents as ascendants."""

    parents = kwargs.get("parents", [])
    child = kwargs.get("child", None)

    assert len(parents) > 0, "Missing input for parents"
    assert child, "Missing input for child"

    for p in parents:
        assert os.access(f"{p}.json", os.F_OK), f"{p}.json does not exist"
    assert os.access(f"{child}.json", os.F_OK), f"{child}.json does not exist"

    for p in parents:
        ascendant = person()
        ascendant.load(p)
        ascendant.data["direct_relatives"]["descendants"].append(child)
        ascendant.save()

        descendant = person()
        descendant.load(child)
        descendant.data["direct_relatives"]["ascendants"].append(p)
        descendant.save()


def link_partners(partners):
    """Add corresponding ids as partners"""

    try:
        assert len(partners) == 2, "Takes two partners only"
        a = person()
        a.load(partners[0])
        a.data["direct_relatives"]["partners"].append(partners[1])
        a.save()
        b = person()
        b.load(partners[1])
        b.data["direct_relatives"]["partners"].append(partners[0])
        b.save()
    except AssertionError as error:
        print(error)


def index():
    """Return a sorted list of all person records."""

    count = 0
    unsorted_id_name = []
    for file in os.listdir("."):
        if len(file) == 11 and file.endswith(".json"):
            count += 1
            p = person()
            p.load(str(file)[:-5])
            unsorted_id_name.append(f"{p.data['id']} - {p.data['last_names'][0]}, {' '.join(p.data['first_names'])} dob {p.data['date_of_birth']}.")

    sorted_names = []
    for name in unsorted_id_name:
        sorted_names.append(name[9:])
    sorted_names.sort()

    sorted_index = []
    for name in sorted_names:
        for id_name in unsorted_id_name:
            if name in id_name:
                if id_name not in sorted_index:
                    sorted_index.append(id_name)
    assert len(sorted_index) == count, "Error sorting index"

    for record in sorted_index:
        print(record)
    print(f"{count} records.")

def orphan_records():
    """Return a list of records which have no relatives linked."""

    orphans = []
    for file in os.listdir("."):
        if len(file) == 11 and file.endswith(".json"):
            count = 0
            p = person()
            p.load(str(file)[:-5])
            for relative in p.data['direct_relatives']:
                count += len(p.data['direct_relatives'][relative])
            if count == 0:
                orphans.append(f"{p.data['id']} - {p.data['last_names'][0]}, {' '.join(p.data['first_names'])} dob {p.data['date_of_birth']}.")

    for orphan in orphans:
        print(orphan)

def find(**kwargs):
    """Scan through .json files and return ids of matches. Query first names, last names, dob, occupation, addresses."""

    records = list()
    for file in os.listdir("."):
        if len(file) == 11 and file.endswith(".json"):
            p = person()
            p.load(file[:-5])
            records.append(p.data)

    results = list()

    first_name = kwargs.get("fn", None)
    if first_name:
        for r in records:
            if first_name in r["first_names"]:
                results.append(r)

    last_name = kwargs.get("ln", None)
    if last_name:
        for r in records:
            if last_name in r["last_names"]:
                results.append(r)

    dob = kwargs.get("dob", None)
    if dob:
        for r in records:
            if dob in r["date_of_birth"]:
                results.append(r)

    occupation = kwargs.get("occupation", None)
    if occupation:
        for r in records:
            for event in r["life_events"]:
                if (
                    event["description"].lower() == "occupation"
                    and occupation.lower() in event["notes"].lower()
                ):
                    results.append(r)

    address = kwargs.get("address", None)
    if address:
        for r in records:
            for event in r["life_events"]:
                if (
                    event["description"].lower() == "address"
                    and address.lower() in event["notes"].lower()
                ):
                    results.append(r)

    for r in results:
        print(
            f"{r['id']} - "
            f"{r['last_names'][0]}, {' '.join(r['first_names'])} "
            f"dob {r['date_of_birth']}."
        )


def request_letter(person, event, output_filename):
    """Request letter for life_event """

    document = docx.Document("templates/Generic Request.docx")
    paragraph = document.add_paragraph(datetime.date.today().strftime("%d/%m/%Y"))
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph("Their address")
    document.add_paragraph("Dear Sir or Ma'am,")
    dob = datetime.datetime.strptime(person.data["date_of_birth"], "%Y-%m-%d").strftime(
        "%d/%m/%Y"
    )
    document.add_paragraph(
        "I am researching my ancestry and to that end submit this request for information to you."
    )
    main_text = document.add_paragraph(
        f"I seek any records you may hold in relation to the {event['description']} of {person.data['first_names'][0]} {person.data['last_names'][0]} dob {dob}. "
    )
    if event["start"]:
        start = datetime.datetime.strptime(event["start"], "%Y-%m-%d").strftime(
            "%d/%m/%Y"
        )
        main_text.add_run(f"The earliest date I have in relation to this is {start}. ")
    if event["finish"]:
        finish = datetime.datetime.strptime(event["finish"], "%Y-%m-%d").strftime(
            "%d/%m/%Y"
        )
        main_text.add_run(f"The last date I have in relation to this is {finish}. ")
    document.add_paragraph(
        "Please let me know if there is a cost I must pay for your help and how I can pay. "
    )
    document.add_paragraph("I am grateful for your assistance. Thank you in advance.")
    document.add_paragraph("Yours sincerely,")
    document.add_paragraph(document.paragraphs[0].text)
    document.save(output_filename)
